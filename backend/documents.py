import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import uuid
from datetime import datetime

# Text extraction imports
import PyPDF2
from docx import Document as DocxDocument
import markdown

from utils import get_file_hash
from schemas import DocumentResponse

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_types = ["pdf", "docx", "txt", "md"]
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from document based on file type"""
        try:
            if file_type == "pdf":
                return self._extract_pdf_text(file_path)
            elif file_type == "docx":
                return self._extract_docx_text(file_path)
            elif file_type == "txt":
                return self._extract_txt_text(file_path)
            elif file_type == "md":
                return self._extract_markdown_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return "\n".join(text)
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_markdown_text(self, file_path: str) -> str:
        """Extract text from Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            # Convert to plain text (remove markdown formatting)
            html = markdown.markdown(md_content)
            import re
            text = re.sub('<[^<]+?>', '', html)
            return text
    
    def chunk_document(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """Chunk document text into smaller pieces"""
        if not text.strip():
            return []
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Use adaptive chunking
        chunks = self._adaptive_chunk(text, chunk_size, overlap)
        
        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk_text in enumerate(chunks):
            chunk_objects.append({
                "id": str(uuid.uuid4()),
                "index": i,
                "text": chunk_text,
                "length": len(chunk_text),
                "created_at": datetime.now()
            })
        
        return chunk_objects
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.,!?;:()\-\'""]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def _adaptive_chunk(self, text: str, target_size: int, overlap: int) -> List[str]:
        """Adaptive chunking that respects sentence boundaries"""
        if len(text) <= target_size:
            return [text]
        
        # Split into sentences
        sentences = self._split_sentences(text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # If adding this sentence would exceed target size
            if len(current_chunk) + len(sentence) > target_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, overlap)
                current_chunk = overlap_text + " " + sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        import re
        
        # Simple sentence splitting (can be improved with NLTK)
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get overlap text from the end of current chunk"""
        if len(text) <= overlap_size:
            return text
        
        # Try to get overlap at word boundary
        overlap_text = text[-overlap_size:]
        space_index = overlap_text.find(' ')
        if space_index > 0:
            return overlap_text[space_index:].strip()
        
        return overlap_text

def process_document_with_embeddings(doc_id: str, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Process chunks and generate embeddings"""
    try:
        from .embeddings import embedding_engine
        
        # Extract text from chunks
        chunk_texts = [chunk["text"] for chunk in chunks]
        
        # Generate embeddings
        logger.info(f"Generating embeddings for {len(chunk_texts)} chunks")
        embeddings = embedding_engine.embed_texts(chunk_texts)
        
        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            chunk["embedding"] = embedding
            chunk["embedding_dim"] = len(embedding)
        
        logger.info(f"Embeddings generated successfully for document: {doc_id}")
        return chunks
        
    except Exception as e:
        logger.error(f"Embedding generation failed: {e}")
        raise

# Database functions
def save_document_to_db(
    db,
    doc_id: str,
    title: str,
    file_path: str,
    file_type: str,
    file_size: int,
    text_preview: str,
    user_id: str = "default_user"
):
    """Save document metadata to database"""
    try:
        from database import Document as DBDocument, User as DBUser
        
        # Create or get user
        user = db.query(DBUser).filter(DBUser.id == user_id).first()
        if not user:
            user = DBUser(
                id=user_id,
                email=f"{user_id}@localhost",
                name="Default User"
            )
            db.add(user)
            db.commit()
        
        file_hash = get_file_hash(file_path)
        
        db_document = DBDocument(
            id=doc_id,
            title=title,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            text_preview=text_preview,
            file_hash=file_hash,
            owner_id=user_id,
            processing_status="extracted"
        )
        
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        return db_document
    except Exception as e:
        logger.error(f"Database save failed: {e}")
        raise

def save_chunks_to_db(db, doc_id: str, chunks: List[Dict[str, Any]]):
    """Save document chunks to database"""
    try:
        from database import DocumentChunk as DBChunk
        
        db_chunks = []
        
        for chunk in chunks:
            db_chunk = DBChunk(
                id=chunk["id"],
                document_id=doc_id,
                chunk_index=chunk["index"],
                chunk_text=chunk["text"],
                chunk_length=chunk["length"],
                embedding_id=chunk["id"],  # Use same ID for Qdrant
                metadata={"created_at": chunk["created_at"].isoformat()}
            )
            db_chunks.append(db_chunk)
        
        db.add_all(db_chunks)
        db.commit()
        
        return db_chunks
    except Exception as e:
        logger.error(f"Chunks save failed: {e}")
        raise

def update_document_status(db, doc_id: str, status: str, chunk_count: int = None):
    """Update document processing status"""
    try:
        from database import Document as DBDocument
        
        document = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
        if document:
            document.processing_status = status
            if chunk_count is not None:
                document.chunk_count = chunk_count
            db.commit()
    except Exception as e:
        logger.error(f"Status update failed: {e}")

# Global processor instance
document_processor = DocumentProcessor()